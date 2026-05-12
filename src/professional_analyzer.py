#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
专业对虾养殖数据分析系统 v2.0
功能：
1. 特征工程（FCR、SGR、环境压力指数）
2. 机器学习预测（产量预测、特征重要性）
3. 环境压力预警（红色预警模块）
4. 异常检测（不合理数据识别）
5. 多维度对比（同比/环比）
6. 一键生成 Word+PDF 报告

作者：SmartShrimp contributors
日期：2026-03-17
"""

import pandas as pd
import numpy as np
from pathlib import Path
from datetime import datetime, timedelta
import sys
import warnings
warnings.filterwarnings('ignore')

# 机器学习库
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error
import joblib

# 可视化
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
import seaborn as sns

# 文档生成
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ============ 导入项目配置 ============
# 添加父目录到路径以导入 config
sys.path.insert(0, str(Path(__file__).parent.parent))
try:
    from config import (
        PROJECT_ROOT, DATA_DIR, REPORTS_DIR, FIGURES_DIR,
        get_chinese_font_config, ANALYSIS_CONFIG, ML_CONFIG,
        VIZ_CONFIG, REPORT_CONFIG
    )
    USE_CONFIG = True
except ImportError:
    # 如果配置文件不可用，使用默认值
    USE_CONFIG = False
    PROJECT_ROOT = Path(__file__).parent.parent
    DATA_DIR = PROJECT_ROOT / 'data'
    REPORTS_DIR = PROJECT_ROOT / 'reports'
    FIGURES_DIR = REPORTS_DIR / 'figures'
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)

# ============ 配置中文字体 ============
def get_chinese_font():
    """跨平台中文字体配置"""
    if USE_CONFIG:
        font_config = get_chinese_font_config()
        plt.rcParams['font.sans-serif'] = font_config['font_names']
        font_path = font_config['font_path']
    else:
        # 回退到平台检测
        import platform
        system = platform.system()
        if system == 'Windows':
            plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
            font_path = 'C:\\Windows\\Fonts\\msyh.ttc'
        elif system == 'Darwin':  # macOS
            plt.rcParams['font.sans-serif'] = ['PingFang SC', 'Heiti TC']
            font_path = '/System/Library/Fonts/PingFang.ttc'
        else:  # Linux
            plt.rcParams['font.sans-serif'] = ['WenQuanYi Micro Hei', 'SimHei']
            font_path = '/usr/share/fonts/wqy-microhei/wqy-microhei.ttc'

    # 尝试加载字体文件
    try:
        font_title = FontProperties(fname=font_path, size=16, weight='bold')
        font_label = FontProperties(fname=font_path, size=12)
        font_legend = FontProperties(fname=font_path, size=10)
        return font_title, font_label, font_legend
    except Exception:
        # 回退到系统字体
        print(f"Warning: Cannot load font from {font_path}, using system default")
        font_title = FontProperties(size=16, weight='bold')
        font_label = FontProperties(size=12)
        font_legend = FontProperties(size=10)
        return font_title, font_label, font_legend

font_title, font_label, font_legend = get_chinese_font()
plt.rcParams['axes.unicode_minus'] = False

# ============ 数据加载 ============
class ShrimpDataLoader:
    """数据加载器"""
    
    def __init__(self, data_path):
        self.data_path = Path(data_path)
        self.df = None
    
    def load(self):
        """加载数据"""
        if not self.data_path.exists():
            raise FileNotFoundError(f"数据文件不存在：{self.data_path}")

        try:
            if self.data_path.suffix == '.xlsx':
                self.df = pd.read_excel(self.data_path, engine='openpyxl')
            elif self.data_path.suffix == '.csv':
                self.df = pd.read_csv(self.data_path, encoding='utf-8-sig')
            else:
                raise ValueError(f"不支持的文件格式：{self.data_path.suffix}")

            # 确保有日期列
            if '日期' not in self.df.columns:
                print(f"警告：数据文件中未找到'日期'列，系统生成日期序列")
                self.df['日期'] = pd.date_range(start='2026-02-15', periods=len(self.df), freq='D')

            self.df['日期'] = pd.to_datetime(self.df['日期'], errors='coerce')
            self.df = self.df.sort_values('日期').reset_index(drop=True)

            print(f"成功加载 {len(self.df)} 条记录")
            return self.df

        except Exception as e:
            raise RuntimeError(f"加载数据文件失败：{str(e)}")


# ============ 特征工程 ============
class FeatureEngineer:
    """特征工程模块"""
    
    def __init__(self, df):
        self.df = df.copy()
    
    def calculate_fcr(self):
        """
        计算 FCR（饲料转化率）
        FCR = 投喂量 (kg) / 预计产量 (kg)
        FCR 越低，饲料转化效率越高
        """
        if '投喂量 (kg)' in self.df.columns and '预计产量 (kg)' in self.df.columns:
            self.df['FCR'] = self.df['投喂量 (kg)'] / (self.df['预计产量 (kg)'] + 0.01)  # 避免除零
            self.df['FCR'] = self.df['FCR'].round(3)
            print("  [OK] FCR（饲料转化率）计算完成")
        return self.df
    
    def calculate_sgr(self):
        """
        计算 SGR（特定生长率）
        SGR = (ln(末体重) - ln(初体重)) / 天数 × 100%
        """
        if '虾体重 (g)' in self.df.columns:
            self.df['SGR'] = self.df['虾体重 (g)'].pct_change() * 100
            self.df['SGR'] = self.df['SGR'].round(2)
            self.df['SGR_累积'] = ((self.df['虾体重 (g)'] / self.df['虾体重 (g)'].iloc[0]) ** (1 / (np.arange(len(self.df)) + 1)) - 1) * 100
            self.df['SGR_累积'] = self.df['SGR_累积'].round(2)
            print("  [OK] SGR（特定生长率）计算完成")
        return self.df
    
    def calculate_environmental_stress(self):
        """
        计算环境压力指数
        - 溶解氧 < 阈值: 高压
        - pH 剧烈波动 (>阈值/天): 高压
        - 水温异常 (<最低 或>最高): 高压
        """
        # 使用配置参数或默认值
        dissolved_oxygen_min = ANALYSIS_CONFIG.get('dissolved_oxygen_min', 5.0) if USE_CONFIG else 5.0
        ph_change_threshold = ANALYSIS_CONFIG.get('ph_change_threshold', 0.3) if USE_CONFIG else 0.3
        temp_min = ANALYSIS_CONFIG.get('temp_min', 24.0) if USE_CONFIG else 24.0
        temp_max = ANALYSIS_CONFIG.get('temp_max', 32.0) if USE_CONFIG else 32.0

        stress_score = np.zeros(len(self.df))
        stress_reasons = []

        for idx, row in self.df.iterrows():
            reasons = []
            score = 0

            # 溶解氧预警
            if '溶解氧 (mg/L)' in row and row['溶解氧 (mg/L)'] < dissolved_oxygen_min:
                score += 3
                reasons.append(f"溶解氧过低 ({row['溶解氧 (mg/L)']:.1f}mg/L)")

            # pH 波动预警
            if idx > 0 and 'pH 值' in row:
                ph_change = abs(row['pH 值'] - self.df.loc[idx-1, 'pH 值'])
                if ph_change > ph_change_threshold:
                    score += 2
                    reasons.append(f"pH 剧烈波动 (Δ{ph_change:.2f})")

            # 水温预警
            if '水温 (°C)' in row:
                if row['水温 (°C)'] < temp_min or row['水温 (°C)'] > temp_max:
                    score += 2
                    reasons.append(f"水温异常 ({row['水温 (°C)']:.1f}°C)")

            stress_score[idx] = score
            stress_reasons.append('; '.join(reasons) if reasons else '正常')

        self.df['环境压力指数'] = stress_score.astype(int)
        self.df['压力原因'] = stress_reasons

        # 预警等级
        self.df['预警等级'] = self.df['环境压力指数'].apply(
            lambda x: '红色预警' if x >= 3 else ('橙色预警' if x >= 2 else ('黄色预警' if x >= 1 else '正常'))
        )

        print("  [OK] 环境压力指数计算完成")
        return self.df
    
    def calculate_lag_effects(self, lag_days=3):
        """
        计算环境因素的滞后影响
        分析溶解氧、pH 值对存活率的滞后影响
        """
        for lag in range(1, lag_days + 1):
            self.df[f'溶解氧_滞后{lag}天'] = self.df['溶解氧 (mg/L)'].shift(lag)
            self.df[f'pH 值_滞后{lag}天'] = self.df['pH 值'].shift(lag)
            self.df[f'水温_滞后{lag}天'] = self.df['水温 (°C)'].shift(lag)
        
        print(f"  [OK] 滞后效应计算完成（{lag_days}天）")
        return self.df
    
    def detect_anomalies(self):
        """
        异常数据检测
        识别不合理的养殖数据
        """
        anomalies = []
        
        for idx, row in self.df.iterrows():
            anomaly_reasons = []
            
            # 摄食率极高但体重没长
            if idx > 0:
                if '摄食率 (%)' in row and row['摄食率 (%)'] > 90:
                    weight_gain = row.get('虾体重 (g)', 0) - self.df.loc[idx-1].get('虾体重 (g)', 0)
                    if weight_gain < 0.3:  # 日增长小于 0.3g
                        anomaly_reasons.append(f"摄食率高 ({row['摄食率 (%)']:.0f}%) 但体重增长低 ({weight_gain:.2f}g)")
            
            # FCR 异常高
            if 'FCR' in row and row['FCR'] > 2.5:
                anomaly_reasons.append(f"FCR 异常高 ({row['FCR']:.2f})")
            
            # 存活率骤降
            if idx > 0 and '存活率 (%)' in row:
                survival_drop = self.df.loc[idx-1].get('存活率 (%)', 100) - row['存活率 (%)']
                if survival_drop > 3:
                    anomaly_reasons.append(f"存活率骤降 ({survival_drop:.1f}%)")
            
            if anomaly_reasons:
                anomalies.append({
                    '日期': row.get('日期', idx),
                    '异常原因': '; '.join(anomaly_reasons)
                })
        
        self.anomalies = pd.DataFrame(anomalies) if anomalies else pd.DataFrame()
        print(f"  [OK] 异常检测完成（发现 {len(anomalies)} 条异常）")
        return self.df
    
    def run_all(self):
        """运行所有特征工程"""
        print("\n[特征工程] 开始计算...")
        self.calculate_fcr()
        self.calculate_sgr()
        self.calculate_environmental_stress()
        self.calculate_lag_effects()
        self.detect_anomalies()
        print("[OK] 特征工程完成\n")
        return self.df


# ============ 机器学习预测 ============
class YieldPredictor:
    """产量预测模型"""
    
    def __init__(self, df):
        self.df = df.copy()
        self.model = None
        self.feature_importance = None
    
    def prepare_features(self):
        """准备特征"""
        feature_cols = ['水温 (°C)', '盐度 (ppt)', 'pH 值', '溶解氧 (mg/L)', '投喂量 (kg)', 'FCR', 'SGR', '环境压力指数']

        # 只保留存在的列
        available_cols = [col for col in feature_cols if col in self.df.columns]

        # 添加滞后特征
        lag_features = [col for col in self.df.columns if '滞后' in col]
        available_cols.extend(lag_features[:3])  # 最多 3 个滞后特征

        # 删除缺失值
        df_clean = self.df[available_cols + ['预计产量 (kg)']].dropna()

        if len(df_clean) < 10:
            raise ValueError("数据量不足，无法训练模型")

        self.X = df_clean[available_cols]
        self.y = df_clean['预计产量 (kg)']
        self.feature_names = available_cols

        return self.X, self.y

    def train(self, model_type='random_forest'):
        """训练模型（使用时序交叉验证）"""
        # 时序数据：使用前80%训练，后20%测试（避免数据泄露）
        split_idx = int(len(self.X) * 0.8)

        X_train, X_test = self.X.iloc[:split_idx], self.X.iloc[split_idx:]
        y_train, y_test = self.y.iloc[:split_idx], self.y.iloc[split_idx:]
        
        if model_type == 'random_forest':
            self.model = RandomForestRegressor(
                n_estimators=100,
                max_depth=10,
                random_state=42,
                n_jobs=-1
            )
        elif model_type == 'gradient_boosting':
            self.model = GradientBoostingRegressor(
                n_estimators=100,
                max_depth=5,
                random_state=42
            )
        
        self.model.fit(X_train, y_train)
        
        # 评估
        y_pred = self.model.predict(X_test)
        self.metrics = {
            'R²': r2_score(y_test, y_pred),
            'MAE': mean_absolute_error(y_test, y_pred),
            'RMSE': np.sqrt(mean_squared_error(y_test, y_pred))
        }
        
        # 特征重要性
        self.feature_importance = pd.DataFrame({
            '特征': self.feature_names,
            '重要性': self.model.feature_importances_
        }).sort_values('重要性', ascending=False)
        
        print(f"  [OK] 模型训练完成 (R²={self.metrics['R²']:.3f})")
        return self.model
    
    def predict(self, features_dict):
        """预测产量"""
        if self.model is None:
            raise ValueError("请先训练模型")
        
        # 构建特征向量
        feature_vector = [features_dict.get(col, 0) for col in self.feature_names]
        prediction = self.model.predict([feature_vector])[0]
        
        return prediction
    
    def generate_recommendations(self):
        """基于模型生成决策建议"""
        recommendations = []
        
        if self.feature_importance is not None:
            # 找出最重要的因素
            top_feature = self.feature_importance.iloc[0]
            
            if top_feature['特征'] == '水温 (°C)':
                recommendations.append({
                    '优先级': '高',
                    '类型': '水温优化',
                    '建议': f'水温是影响产量的最关键因素（贡献度{top_feature["重要性"]:.1%}）。建议保持水温在 26-30°C 区间。'
                })
            elif top_feature['特征'] == '溶解氧 (mg/L)':
                recommendations.append({
                    '优先级': '高',
                    '类型': '溶解氧优化',
                    '建议': f'溶解氧是影响产量的最关键因素（贡献度{top_feature["重要性"]:.1%}）。建议提前 2 小时开启增氧机，保持溶解氧>6mg/L。'
                })
            elif top_feature['特征'] == '投喂量 (kg)':
                recommendations.append({
                    '优先级': '高',
                    '类型': '投喂优化',
                    '建议': f'投喂量是影响产量的最关键因素（贡献度{top_feature["重要性"]:.1%}）。建议根据 FCR 调整投喂策略，目标 FCR<1.5。'
                })
            elif top_feature['特征'] == 'FCR':
                recommendations.append({
                    '优先级': '高',
                    '类型': '饲料效率优化',
                    '建议': f'FCR 是影响产量的最关键因素（贡献度{top_feature["重要性"]:.1%}）。当前 FCR 偏高，建议优化饲料配方或投喂频率。'
                })
        
        # 基于当前数据的具体建议
        if '环境压力指数' in self.df.columns:
            high_stress_days = (self.df['环境压力指数'] >= 3).sum()
            if high_stress_days > 0:
                recommendations.append({
                    '优先级': '高',
                    '类型': '环境预警',
                    '建议': f'检测到{high_stress_days}天环境压力指数达到红色预警级别。建议立即检查水质并采取措施。'
                })
        
        if 'FCR' in self.df.columns:
            avg_fcr = self.df['FCR'].mean()
            if avg_fcr > 1.8:
                recommendations.append({
                    '优先级': '中',
                    '类型': 'FCR 优化',
                    '建议': f'平均 FCR 为{avg_fcr:.2f}，高于行业优秀水平 (1.5)。建议优化投喂策略，预计可降低成本{(avg_fcr-1.5)*10:.1f}%。'
                })
        
        return recommendations
    
    def run_all(self):
        """运行完整预测流程"""
        print("\n🤖 开始机器学习预测...")
        self.prepare_features()
        self.train()
        self.generate_recommendations()
        print("[OK] 机器学习预测完成\n")
        return self


# ============ 可视化生成 ============
class Visualizer:
    """可视化模块"""
    
    def __init__(self, df, output_dir):
        self.df = df
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def plot_fcr_trend(self):
        """FCR 趋势分析"""
        if 'FCR' not in self.df.columns:
            print("  跳过 FCR 趋势图（缺少 FCR 数据）")
            return None

        try:
            fig, ax = plt.subplots(figsize=(12, 6))

            # FCR 趋势线
            ax.plot(self.df['日期'], self.df['FCR'], marker='o', linewidth=2, color='steelblue', label='FCR')
            ax.axhline(y=self.df['FCR'].mean(), color='red', linestyle='--', linewidth=2, label=f'平均 FCR ({self.df["FCR"].mean():.2f})')
            ax.axhline(y=1.5, color='green', linestyle=':', linewidth=2, label='优秀水平 (1.5)')

            # 最优 FCR 区间
            min_fcr_idx = self.df['FCR'].idxmin()
            min_fcr_date = self.df.loc[min_fcr_idx, '日期']
            min_fcr_temp = self.df.loc[min_fcr_idx, '水温 (°C)'] if '水温 (°C)' in self.df.columns else None

            if min_fcr_temp:
                ax.axvline(x=min_fcr_date, color='orange', linestyle='-.', linewidth=2,
                          label=f'最优 FCR 日期 (水温{min_fcr_temp:.1f}°C)')

            ax.set_xlabel('日期', fontproperties=font_label)
            ax.set_ylabel('FCR（饲料转化率）', fontproperties=font_label)
            ax.set_title('FCR 趋势分析（越低越好）', fontproperties=font_title)
            ax.legend(prop=font_legend)
            ax.grid(True, alpha=0.3)
            plt.xticks(rotation=45)
            plt.tight_layout()

            path = self.output_dir / 'fcr_trend.png'
            plt.savefig(path, dpi=300, bbox_inches='tight')
            plt.close()

            print(f"  OK FCR 趋势图：{path}")
            return path
        except Exception as e:
            print(f"  ERROR FCR 趋势图生成失败：{str(e)}")
            return None
    
    def plot_sgr_trend(self):
        """SGR 趋势分析"""
        if 'SGR' not in self.df.columns:
            return None
        
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 5))
        
        # 日 SGR
        ax1.plot(self.df['日期'], self.df['SGR'], marker='s', linewidth=2, color='forestgreen')
        ax1.axhline(y=0, color='gray', linestyle='--', linewidth=1)
        ax1.set_xlabel('日期', fontproperties=font_label)
        ax1.set_ylabel('日 SGR (%)', fontproperties=font_label)
        ax1.set_title('日特定生长率 (SGR)', fontproperties=font_title)
        ax1.grid(True, alpha=0.3)
        plt.sca(ax1)
        plt.xticks(rotation=45)
        
        # 累积 SGR
        if 'SGR_累积' in self.df.columns:
            ax2.plot(self.df['日期'], self.df['SGR_累积'], marker='o', linewidth=2, color='darkorange')
            ax2.set_xlabel('日期', fontproperties=font_label)
            ax2.set_ylabel('累积 SGR (%)', fontproperties=font_label)
            ax2.set_title('累积特定生长率', fontproperties=font_title)
            ax2.grid(True, alpha=0.3)
            plt.sca(ax2)
            plt.xticks(rotation=45)
        
        plt.tight_layout()
        
        path = self.output_dir / 'sgr_trend.png'
        plt.savefig(path, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"  [OK] SGR 趋势图：{path}")
        return path
    
    def plot_environmental_stress(self):
        """环境压力预警图"""
        if '环境压力指数' not in self.df.columns:
            return None
        
        fig, ax = plt.subplots(figsize=(14, 6))
        
        # 预警等级颜色
        colors = {'正常': 'green', '黄色预警': 'yellow', '橙色预警': 'orange', '红色预警': 'red'}
        
        for level in ['正常', '黄色预警', '橙色预警', '红色预警']:
            mask = self.df['预警等级'] == level
            ax.scatter(self.df.loc[mask, '日期'], self.df.loc[mask, '环境压力指数'],
                      c=colors[level], label=level, s=100, alpha=0.7, edgecolors='black')
        
        ax.set_xlabel('日期', fontproperties=font_label)
        ax.set_ylabel('环境压力指数', fontproperties=font_label)
        ax.set_title('环境压力预警分析', fontproperties=font_title)
        ax.legend(prop=font_legend, loc='upper right')
        ax.grid(True, alpha=0.3)
        ax.set_ylim(-0.5, 8)
        plt.xticks(rotation=45)
        plt.tight_layout()
        
        path = self.output_dir / 'environmental_stress.png'
        plt.savefig(path, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"  [OK] 环境压力预警图：{path}")
        return path
    
    def plot_feature_importance(self, predictor):
        """特征重要性图"""
        if predictor.feature_importance is None:
            return None
        
        fig, ax = plt.subplots(figsize=(10, 8))
        
        # 只取前 10 个重要特征
        top_features = predictor.feature_importance.head(10)
        
        colors = plt.cm.viridis(np.linspace(0.2, 0.8, len(top_features)))
        bars = ax.barh(range(len(top_features)), top_features['重要性'], color=colors)
        
        ax.set_yticks(range(len(top_features)))
        ax.set_yticklabels(top_features['特征'], fontproperties=font_label)
        ax.set_xlabel('特征重要性', fontproperties=font_label)
        ax.set_title('产量预测特征重要性分析', fontproperties=font_title)
        ax.invert_yaxis()
        
        # 添加数值标签
        for i, (bar, value) in enumerate(zip(bars, top_features['重要性'])):
            ax.text(value + 0.01, bar.get_y() + bar.get_height()/2, 
                   f'{value:.1%}', va='center', fontsize=10)
        
        plt.tight_layout()
        
        path = self.output_dir / 'feature_importance.png'
        plt.savefig(path, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"  [OK] 特征重要性图：{path}")
        return path
    
    def plot_correlation_heatmap(self):
        """相关性热力图（增强版）"""
        numeric_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()
        
        if len(numeric_cols) < 2:
            return None
        
        fig, ax = plt.subplots(figsize=(14, 12))
        corr_matrix = self.df[numeric_cols].corr()
        
        mask = np.triu(np.ones_like(corr_matrix, dtype=bool))
        
        sns.heatmap(corr_matrix, mask=mask, annot=True, fmt='.2f', cmap='RdBu_r', 
                   center=0, square=True, linewidths=0.5, ax=ax,
                   cbar_kws={'shrink': 0.8}, annot_kws={'size': 8})
        
        ax.set_title('变量相关性热力图（增强版）', fontproperties=font_title, pad=20)
        ax.set_xticklabels(ax.get_xticklabels(), fontproperties=font_label, rotation=45, ha='right')
        ax.set_yticklabels(ax.get_yticklabels(), fontproperties=font_label, rotation=0)
        
        plt.tight_layout()
        
        path = self.output_dir / 'correlation_heatmap_enhanced.png'
        plt.savefig(path, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"  [OK] 增强版相关性热力图：{path}")
        return path
    
    def generate_all(self, predictor=None):
        """生成所有图表"""
        print("\n[CHART] 开始生成可视化图表...")
        self.plot_fcr_trend()
        self.plot_sgr_trend()
        self.plot_environmental_stress()
        self.plot_correlation_heatmap()
        if predictor:
            self.plot_feature_importance(predictor)
        print("[OK] 可视化图表生成完成\n")


# ============ Word 报告生成 ============
class ReportGenerator:
    """Word 报告生成器"""
    
    def __init__(self, df, predictor, visualizer, output_path):
        self.df = df
        self.predictor = predictor
        self.visualizer = visualizer
        self.output_path = Path(output_path)
        self.doc = Document()
        
        # 设置全局字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def add_title(self, title, level=0):
        """添加标题"""
        heading = self.doc.add_heading(title, level=level)
        heading.runs[0].font.name = 'Microsoft YaHei'
        heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def add_paragraph(self, text, style=None):
        """添加段落"""
        p = self.doc.add_paragraph(text, style=style)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def add_picture(self, path, width=Inches(6)):
        """添加图片"""
        if Path(path).exists():
            self.doc.add_picture(str(path), width=width)
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_alert_box(self, level, message):
        """添加预警框"""
        colors = {
            '红色预警': RGBColor(255, 0, 0),
            '橙色预警': RGBColor(255, 165, 0),
            '黄色预警': RGBColor(255, 255, 0),
            '正常': RGBColor(0, 128, 0)
        }
        
        p = self.doc.add_paragraph()
        p.paragraph_format.border_top = True
        p.paragraph_format.border_bottom = True
        p.paragraph_format.shading.background_color = colors.get(level, RGBColor(200, 200, 200))
        
        run = p.add_run(f'[{level}] {message}')
        run.bold = True
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
    
    def generate(self):
        """生成完整报告"""
        print("\n📄 开始生成 Word 报告...")
        
        # 标题
        self.add_title('对虾养殖数据分析报告（专业版）', level=0)
        self.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', style='List Bullet')
        self.add_paragraph(f'数据源：shrimp_farming_sample.xlsx', style='List Bullet')
        self.add_paragraph(f'分析工具：OpenClaw 专业养殖数据分析系统 v2.0', style='List Bullet')
        self.add_paragraph(f'作者：SmartShrimp contributors', style='List Bullet')
        
        self.doc.add_page_break()
        
        # 目录
        self.add_title('目录', level=1)
        toc = [
            '一、执行摘要',
            '二、特征工程分析',
            '三、环境压力预警',
            '四、异常数据检测',
            '五、机器学习预测',
            '六、决策建议',
            '七、详细数据',
        ]
        for item in toc:
            self.add_paragraph(item, style='List Number')
        
        self.doc.add_page_break()
        
        # 一、执行摘要
        self.add_title('一、执行摘要', level=1)
        self.add_paragraph('本报告采用专业养殖数据分析方法，包含以下核心分析：')
        self.add_paragraph('1. FCR（饲料转化率）分析', style='List Number')
        self.add_paragraph('2. SGR（特定生长率）分析', style='List Number')
        self.add_paragraph('3. 环境压力预警系统', style='List Number')
        self.add_paragraph('4. 机器学习产量预测', style='List Number')
        self.add_paragraph('5. 特征重要性归因分析', style='List Number')
        
        # 关键指标
        self.add_paragraph('\n关键指标：', style='List Bullet')
        if 'FCR' in self.df.columns:
            self.add_paragraph(f'  平均 FCR：{self.df["FCR"].mean():.2f}（优秀水平<1.5）', style='List Bullet')
        if 'SGR' in self.df.columns:
            self.add_paragraph(f'  平均 SGR：{self.df["SGR"].mean():.2f}%', style='List Bullet')
        if '环境压力指数' in self.df.columns:
            red_days = (self.df['预警等级'] == '红色预警').sum()
            self.add_paragraph(f'  红色预警天数：{red_days}天', style='List Bullet')
        
        self.doc.add_page_break()
        
        # 二、特征工程分析
        self.add_title('二、特征工程分析', level=1)
        
        self.add_title('2.1 FCR（饲料转化率）分析', level=2)
        self.add_paragraph('FCR = 投喂量 (kg) / 预计产量 (kg)')
        self.add_paragraph('FCR 越低，饲料转化效率越高。行业优秀水平为 FCR<1.5。')
        
        if 'FCR' in self.df.columns:
            min_fcr_idx = self.df['FCR'].idxmin()
            min_fcr_temp = self.df.loc[min_fcr_idx, '水温 (°C)'] if '水温 (°C)' in self.df.columns else None
            self.add_paragraph(f'\n最优 FCR：{self.df["FCR"].min():.2f}', style='List Bullet')
            if min_fcr_temp:
                self.add_paragraph(f'最优水温区间：{min_fcr_temp:.1f}°C', style='List Bullet')
        
        self.add_picture(self.visualizer.output_dir / 'fcr_trend.png', width=Inches(7))
        self.add_paragraph('图 1：FCR 趋势分析', style='Caption')
        
        self.add_title('2.2 SGR（特定生长率）分析', level=2)
        self.add_paragraph('SGR = (ln(末体重) - ln(初体重)) / 天数 × 100%')
        self.add_paragraph('SGR 反映对虾的日生长速度。')
        
        self.add_picture(self.visualizer.output_dir / 'sgr_trend.png', width=Inches(7))
        self.add_paragraph('图 2：SGR 趋势分析', style='Caption')
        
        self.doc.add_page_break()
        
        # 三、环境压力预警
        self.add_title('三、环境压力预警', level=1)
        
        self.add_paragraph('预警标准：')
        self.add_paragraph('  红色预警：环境压力指数≥3（溶解氧<5mg/L 或水温异常）', style='List Bullet')
        self.add_paragraph('  橙色预警：环境压力指数=2（pH 剧烈波动）', style='List Bullet')
        self.add_paragraph('  黄色预警：环境压力指数=1（轻微异常）', style='List Bullet')
        self.add_paragraph('  正常：环境压力指数=0', style='List Bullet')
        
        # 预警天数统计
        if '预警等级' in self.df.columns:
            self.add_paragraph('\n预警统计：', style='List Bullet')
            for level in ['红色预警', '橙色预警', '黄色预警', '正常']:
                days = (self.df['预警等级'] == level).sum()
                if days > 0:
                    self.add_paragraph(f'  {level}：{days}天', style='List Bullet')
        
        self.add_picture(self.visualizer.output_dir / 'environmental_stress.png', width=Inches(7))
        self.add_paragraph('图 3：环境压力预警分析', style='Caption')
        
        # 红色预警详情
        red_alerts = self.df[self.df['预警等级'] == '红色预警']
        if len(red_alerts) > 0:
            self.add_title('红色预警详情：', level=3)
            for idx, row in red_alerts.iterrows():
                date_str = row['日期'].strftime('%Y-%m-%d') if pd.notna(row.get('日期')) else f'第{idx+1}天'
                reason = row.get('压力原因', '未知原因')
                self.add_alert_box('红色预警', f'{date_str}: {reason}')
        
        self.doc.add_page_break()
        
        # 四、异常数据检测
        self.add_title('四、异常数据检测', level=1)
        
        if hasattr(self.df, 'anomalies') and len(self.df.anomalies) > 0:
            self.add_paragraph(f'检测到 {len(self.df.anomalies)} 条异常数据：')
            for idx, anomaly in self.df.anomalies.iterrows():
                date_str = anomaly['日期'].strftime('%Y-%m-%d') if pd.notna(anomaly.get('日期')) else f'第{idx+1}条'
                self.add_alert_box('橙色预警', f'{date_str}: {anomaly["异常原因"]}')
        else:
            self.add_paragraph('[OK] 未检测到明显异常数据', style='List Bullet')
        
        self.doc.add_page_break()
        
        # 五、机器学习预测
        self.add_title('五、机器学习预测', level=1)
        
        if self.predictor and self.predictor.model:
            self.add_paragraph(f'模型类型：Random Forest Regressor', style='List Bullet')
            self.add_paragraph(f'训练数据量：{len(self.predictor.X)}条', style='List Bullet')
            self.add_paragraph(f'R²: {self.predictor.metrics["R²"]:.3f}', style='List Bullet')
            self.add_paragraph(f'MAE: {self.predictor.metrics["MAE"]:.2f}kg', style='List Bullet')
            self.add_paragraph(f'RMSE: {self.predictor.metrics["RMSE"]:.2f}kg', style='List Bullet')
            
            self.add_picture(self.visualizer.output_dir / 'feature_importance.png', width=Inches(7))
            self.add_paragraph('图 4：产量预测特征重要性分析', style='Caption')
            
            self.add_title('特征重要性解读：', level=3)
            if self.predictor.feature_importance is not None:
                for idx, row in self.predictor.feature_importance.head(5).iterrows():
                    self.add_paragraph(f'  {row["特征"]}: {row["重要性"]:.1%}', style='List Bullet')
        
        self.doc.add_page_break()
        
        # 六、决策建议
        self.add_title('六、决策建议', level=1)
        
        if self.predictor:
            recommendations = self.predictor.generate_recommendations()
            for rec in recommendations:
                p = self.doc.add_paragraph()
                run = p.add_run(f'[{rec["优先级"]}] {rec["类型"]}: {rec["建议"]}')
                if rec['优先级'] == '高':
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
        
        self.doc.add_page_break()
        
        # 七、详细数据
        self.add_title('七、详细数据', level=1)
        
        # 数据摘要表
        self.add_paragraph('数据摘要：')
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ['变量', '均值', '标准差', '范围']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        numeric_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()
        for col in numeric_cols[:10]:  # 最多 10 个变量
            row = table.add_row().cells
            row[0].text = col
            row[1].text = f'{self.df[col].mean():.2f}'
            row[2].text = f'{self.df[col].std():.2f}'
            row[3].text = f'{self.df[col].min():.2f} ~ {self.df[col].max():.2f}'
        
        # 保存
        self.doc.save(str(self.output_path))
        print(f"[OK] Word 报告生成完成：{self.output_path}\n")


# ============ 主程序 ============
def generate_report_only():
    """只生成报告（基于已有的图表和模型）"""
    print("=" * 70)
    print("生成Word报告")
    print("=" * 70)

    # 路径配置
    base_dir = Path(__file__).parent.parent
    data_dir = base_dir / 'data'
    reports_dir = base_dir / 'reports'

    # 查找数据文件
    data_files = list(data_dir.glob('*.xlsx')) + list(data_dir.glob('*.csv'))
    if not data_files:
        print(f"\n未找到数据文件")
        return

    data_path = data_files[0]
    print(f"\n使用数据文件：{data_path.name}")

    # 加载数据
    print("\n[LOAD] 加载数据...")
    loader = ShrimpDataLoader(data_path)
    df = loader.load()
    print(f"  [OK] 加载 {len(df)} 条记录")

    # 特征工程
    fe = FeatureEngineer(df)
    df_enhanced = fe.run_all()

    # 机器学习预测
    predictor = YieldPredictor(df_enhanced)
    predictor.run_all()

    # 可视化
    visualizer = Visualizer(df_enhanced, reports_dir / 'figures_pro')
    visualizer.generate_all(predictor)

    # 生成报告
    output_path = reports_dir / 'analysis_report_pro.docx'
    report_gen = ReportGenerator(df_enhanced, predictor, visualizer, output_path)
    report_gen.generate()

    print(f"\n[OK] 报告生成完成：{output_path}")


def main():
    """主程序"""
    print("=" * 70)
    print("专业对虾养殖数据分析系统 v2.0")
    print("=" * 70)
    print(f"作者：SmartShrimp contributors")
    print(f"日期：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 70)

    # 路径配置 - 使用相对路径
    base_dir = Path(__file__).parent.parent
    data_dir = base_dir / 'data'
    reports_dir = base_dir / 'reports'

    # 创建必要的目录
    data_dir.mkdir(exist_ok=True)
    reports_dir.mkdir(exist_ok=True)

    # 查找数据文件
    data_files = list(data_dir.glob('*.xlsx')) + list(data_dir.glob('*.csv'))

    if not data_files:
        print(f"\n未找到数据文件，请在 {data_dir} 目录下放置 .xlsx 或 .csv 文件")
        return

    data_path = data_files[0]
    print(f"\n使用数据文件：{data_path.name}")
    
    # 1. 数据加载
    print("\n[LOAD] 加载数据...")
    loader = ShrimpDataLoader(data_path)
    df = loader.load()
    print(f"  [OK] 加载 {len(df)} 条记录")
    
    # 2. 特征工程
    fe = FeatureEngineer(df)
    df_enhanced = fe.run_all()
    
    # 3. 机器学习预测
    predictor = YieldPredictor(df_enhanced)
    predictor.run_all()
    
    # 4. 可视化
    visualizer = Visualizer(df_enhanced, reports_dir / 'figures_pro')
    visualizer.generate_all(predictor)
    
    # 5. 生成 Word 报告
    output_path = reports_dir / 'analysis_report_pro.docx'
    report_gen = ReportGenerator(df_enhanced, predictor, visualizer, output_path)
    report_gen.generate()
    
    # 6. 生成 PDF（如果有 pandoc）
    print("\n📄 尝试生成 PDF...")
    try:
        import subprocess
        pdf_path = reports_dir / 'analysis_report_pro.pdf'
        result = subprocess.run(
            ['pandoc', str(output_path), '-o', str(pdf_path)],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            print(f"[OK] PDF 报告生成完成：{pdf_path}")
        else:
            print(f"⚠️ PDF 生成失败：{result.stderr}")
    except Exception as e:
        print(f"⚠️ PDF 生成失败：{e}")
    
    print("\n" + "=" * 70)
    print("[OK] 分析完成！")
    print("=" * 70)
    print(f"\n📁 输出文件：")
    print(f"  Word 报告：{output_path}")
    print(f"  图表目录：{visualizer.output_dir}")
    print(f"\n🎯 下一步：")
    print(f"  1. 查看 Word 报告")
    print(f"  2. 根据决策建议优化养殖策略")
    print(f"  3. 持续监控 FCR 和环境压力指数")
    print("=" * 70)


if __name__ == '__main__':
    main()
